"""
Enterprise extension services for document, workflow, compliance, AI, and SaaS modules.
"""

from datetime import datetime, date, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional
import os
import re
import statistics
import shutil
import uuid

from fastapi import UploadFile
try:
    import requests  # optional; for government integration HTTP calls
    HAS_REQUESTS = True
except ImportError:
    HAS_REQUESTS = False
    requests = None  # type: ignore
from sqlalchemy import and_, func, or_
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.models import (
    AIConversation,
    AIPrediction,
    AISuggestion,
    AIUsageLog,
    ApprovalWorkflow,
    ComplianceNotice,
    ComplianceTask,
    Customer,
    DeviceRegistration,
    DirectorKYC,
    Document,
    DocumentFolder,
    DocumentTag,
    DocumentVersion,
    EmbeddingRecord,
    ExtractionAuditLog,
    FeatureFlag,
    FinancialMetric,
    FraudDetection,
    GovernmentIntegration,
    IntegrationSyncLog,
    Invoice,
    KnowledgeBase,
    Ledger,
    PermissionMatrix,
    PromptTemplate,
    RocFiling,
    SaaSPlan,
    SecurityLog,
    SystemHealthMetric,
    TenantSubscription,
    Transaction,
    UsageMetric,
    UserSession,
    Voucher,
    VoiceCommand,
    WebhookEvent,
    WhatsAppMessage,
    WorkflowHistory,
    WorkflowInstance,
    WorkflowNotification,
    WorkflowStage,
)


UPLOAD_ROOT = Path("uploads")


class DocumentService:
    ALLOWED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls", ".jpg", ".jpeg", ".png", ".zip"}
    ACTIONS = {"view", "upload", "download", "delete", "archive", "manage"}

    @staticmethod
    def _metadata(document: Document) -> Dict[str, Any]:
        return document.extracted_data or {}

    @classmethod
    def has_permission(cls, db: Session, company_id: int, role_name: str, action: str) -> bool:
        if action not in cls.ACTIONS:
            return False
        matrix = db.query(PermissionMatrix).filter(
            PermissionMatrix.company_id == company_id,
            PermissionMatrix.role_name == role_name,
            PermissionMatrix.module_name == "documents",
            PermissionMatrix.is_active == True,
        ).first()
        if not matrix:
            return role_name.lower() in {"admin", "owner"}
        permissions = matrix.permissions or {}
        return bool(permissions.get(action) or permissions.get("manage"))

    @classmethod
    def create_folder(cls, db: Session, company_id: int, folder_name: str, parent_id: Optional[int] = None, description: Optional[str] = None):
        if parent_id:
            parent = db.query(DocumentFolder).filter(
                DocumentFolder.id == parent_id,
                DocumentFolder.company_id == company_id,
            ).first()
            if not parent:
                raise ValueError("Parent folder not found")
        folder = DocumentFolder(
            company_id=company_id,
            parent_id=parent_id,
            folder_name=folder_name,
            description=description,
        )
        db.add(folder)
        db.commit()
        db.refresh(folder)
        return folder

    @classmethod
    def upload_document(
        cls,
        db: Session,
        company_id: int,
        file: UploadFile,
        document_type: str,
        uploaded_by: Optional[str] = None,
        tags: Optional[List[str]] = None,
        folder_id: Optional[int] = None,
        category: Optional[str] = None,
    ):
        if folder_id:
            folder = db.query(DocumentFolder).filter(
                DocumentFolder.id == folder_id,
                DocumentFolder.company_id == company_id,
            ).first()
            if not folder:
                raise ValueError("Folder not found")

        extension = Path(file.filename or "").suffix.lower()
        if extension not in cls.ALLOWED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}")

        upload_dir = UPLOAD_ROOT / str(company_id)
        upload_dir.mkdir(parents=True, exist_ok=True)
        stored_name = f"{uuid.uuid4().hex}{extension}"
        file_path = upload_dir / stored_name

        with file_path.open("wb") as target:
            shutil.copyfileobj(file.file, target)

        metadata = {"folder_id": folder_id, "category": category}
        document = Document(
            company_id=company_id,
            document_type=document_type,
            file_name=file.filename or stored_name,
            file_path=str(file_path),
            file_size=file_path.stat().st_size,
            mime_type=file.content_type,
            tags=tags or [],
            extracted_data=metadata,
            uploaded_by=uploaded_by,
        )
        db.add(document)
        db.flush()

        version = DocumentVersion(
            document_id=document.id,
            version_number=1,
            file_name=document.file_name,
            file_path=document.file_path,
            file_size=document.file_size,
            mime_type=document.mime_type,
            change_notes="Initial upload",
            uploaded_by=uploaded_by,
        )
        db.add(version)
        db.commit()
        db.refresh(document)
        return document

    @classmethod
    def add_version(cls, db: Session, document_id: int, file: UploadFile, uploaded_by: Optional[str] = None, change_notes: Optional[str] = None):
        document = db.query(Document).get(document_id)
        if not document:
            raise ValueError("Document not found")

        extension = Path(file.filename or "").suffix.lower()
        if extension not in cls.ALLOWED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}")

        upload_dir = UPLOAD_ROOT / str(document.company_id)
        upload_dir.mkdir(parents=True, exist_ok=True)
        stored_name = f"{uuid.uuid4().hex}{extension}"
        file_path = upload_dir / stored_name
        with file_path.open("wb") as target:
            shutil.copyfileobj(file.file, target)

        next_version = db.query(DocumentVersion).filter(DocumentVersion.document_id == document_id).count() + 1
        version = DocumentVersion(
            document_id=document_id,
            version_number=next_version,
            file_name=file.filename or stored_name,
            file_path=str(file_path),
            file_size=file_path.stat().st_size,
            mime_type=file.content_type,
            change_notes=change_notes,
            uploaded_by=uploaded_by,
        )
        document.file_name = version.file_name
        document.file_path = version.file_path
        document.file_size = version.file_size
        document.mime_type = version.mime_type
        db.add(version)
        db.commit()
        db.refresh(version)
        return version

    @classmethod
    def delete_document(cls, db: Session, document_id: int):
        document = db.query(Document).get(document_id)
        if not document:
            raise ValueError("Document not found")
        document.is_archived = True
        db.commit()
        return document

    @classmethod
    def hard_delete_document(cls, db: Session, document_id: int):
        document = db.query(Document).get(document_id)
        if not document:
            raise ValueError("Document not found")
        versions = db.query(DocumentVersion).filter(DocumentVersion.document_id == document_id).all()
        paths = [document.file_path] + [version.file_path for version in versions]
        for file_path in paths:
            try:
                path = Path(file_path)
                if path.exists() and UPLOAD_ROOT.resolve() in path.resolve().parents:
                    path.unlink()
            except OSError:
                pass
        for version in versions:
            db.delete(version)
        db.delete(document)
        db.commit()
        return {"deleted": True, "document_id": document_id}

    @classmethod
    def create_tag(cls, db: Session, company_id: int, tag_name: str, tag_color: Optional[str] = None):
        tag = DocumentTag(company_id=company_id, tag_name=tag_name, tag_color=tag_color)
        db.add(tag)
        db.commit()
        db.refresh(tag)
        return tag

    @classmethod
    def search(
        cls,
        db: Session,
        company_id: int,
        query: Optional[str] = None,
        document_type: Optional[str] = None,
        tag: Optional[str] = None,
        folder_id: Optional[int] = None,
        category: Optional[str] = None,
        include_archived: bool = False,
    ):
        documents = db.query(Document).filter(Document.company_id == company_id)
        if not include_archived:
            documents = documents.filter(Document.is_archived == False)
        if document_type:
            documents = documents.filter(Document.document_type == document_type)
        if query:
            like = f"%{query}%"
            documents = documents.filter(or_(
                Document.file_name.ilike(like),
                Document.document_type.ilike(like),
                Document.document_number.ilike(like),
            ))
        result = []
        for document in documents.order_by(Document.uploaded_on.desc()).all():
            metadata = cls._metadata(document)
            tags = document.tags or []
            if tag and tag not in tags:
                continue
            if folder_id is not None and metadata.get("folder_id") != folder_id:
                continue
            if category and metadata.get("category") != category:
                continue
            result.append(document)
        return result


class OCRService:
    GSTIN_PATTERN = re.compile(r"\b\d{2}[A-Z]{5}\d{4}[A-Z][A-Z\d]Z[A-Z\d]\b")
    PAN_PATTERN = re.compile(r"\b[A-Z]{5}\d{4}[A-Z]\b")
    TAN_PATTERN = re.compile(r"\b[A-Z]{4}\d{5}[A-Z]\b")
    INVOICE_PATTERN = re.compile(r"(?:invoice\s*(?:no|number|#)?|inv\s*(?:no|#)?)[:\-\s]*([A-Z0-9\/\-]+)", re.IGNORECASE)
    AMOUNT_PATTERN = re.compile(r"(?:grand\s*total|total\s*amount|amount\s*payable|total)[:\-\s]*(?:rs\.?|inr)?\s*([0-9,]+(?:\.\d{1,2})?)", re.IGNORECASE)
    TAX_PATTERN = re.compile(r"\b(CGST|SGST|IGST|GST|TAX)\b[:\-\s]*(?:rs\.?|inr)?\s*([0-9,]+(?:\.\d{1,2})?)", re.IGNORECASE)

    @classmethod
    def _read_text(cls, document: Document, provider: str) -> str:
        path = Path(document.file_path)
        if not path.exists():
            raise ValueError("Document file not found on disk")

        suffix = path.suffix.lower()
        if provider.lower() in {"tesseract", "local"}:
            return cls._tesseract_text(path, suffix)
        if provider.lower() == "aws":
            return cls._aws_textract_text(path)
        if provider.lower() == "google":
            return cls._google_vision_text(path)
        if provider.lower() == "azure":
            return cls._azure_ocr_text(path)
        raise ValueError(f"Unsupported OCR provider: {provider}")

    @staticmethod
    def _tesseract_text(path: Path, suffix: str) -> str:
        if suffix in {".txt", ".csv"}:
            return path.read_text(encoding="utf-8", errors="ignore")
        if suffix in {".xlsx", ".xls"}:
            try:
                from openpyxl import load_workbook
            except ImportError as exc:
                raise ValueError("openpyxl is required for Excel OCR extraction") from exc
            workbook = load_workbook(path, read_only=True, data_only=True)
            rows = []
            for sheet in workbook.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    rows.append(" ".join(str(cell) for cell in row if cell is not None))
            return "\n".join(rows)
        if suffix == ".docx":
            try:
                from docx import Document as DocxDocument
            except ImportError as exc:
                raise ValueError("python-docx is required for DOCX text extraction") from exc
            return "\n".join(paragraph.text for paragraph in DocxDocument(path).paragraphs)
        if suffix == ".pdf":
            try:
                from pdf2image import convert_from_path
                import pytesseract
            except ImportError as exc:
                raise ValueError("pdf2image and pytesseract are required for PDF OCR") from exc
            text = []
            for page in convert_from_path(str(path)):
                text.append(pytesseract.image_to_string(page))
            return "\n".join(text)
        if suffix in {".jpg", ".jpeg", ".png"}:
            try:
                from PIL import Image
                import pytesseract
            except ImportError as exc:
                raise ValueError("pillow and pytesseract are required for image OCR") from exc
            return pytesseract.image_to_string(Image.open(path))
        if suffix == ".zip":
            raise ValueError("ZIP archives must be extracted before OCR processing")
        raise ValueError(f"OCR is not supported for {suffix}")

    @staticmethod
    def _aws_textract_text(path: Path) -> str:
        try:
            import boto3
        except ImportError as exc:
            raise ValueError("boto3 is required for AWS Textract") from exc
        client = boto3.client("textract")
        response = client.detect_document_text(Document={"Bytes": path.read_bytes()})
        return "\n".join(block.get("Text", "") for block in response.get("Blocks", []) if block.get("BlockType") == "LINE")

    @staticmethod
    def _google_vision_text(path: Path) -> str:
        endpoint = os.getenv("GOOGLE_VISION_ENDPOINT", "https://vision.googleapis.com/v1/images:annotate")
        api_key = os.getenv("GOOGLE_VISION_API_KEY")
        if not api_key:
            raise ValueError("GOOGLE_VISION_API_KEY is required")
        import base64
        payload = {
            "requests": [{
                "image": {"content": base64.b64encode(path.read_bytes()).decode("ascii")},
                "features": [{"type": "DOCUMENT_TEXT_DETECTION"}],
            }]
        }
        response = requests.post(f"{endpoint}?key={api_key}", json=payload, timeout=60)
        response.raise_for_status()
        data = response.json()
        return data["responses"][0].get("fullTextAnnotation", {}).get("text", "")

    @staticmethod
    def _azure_ocr_text(path: Path) -> str:
        endpoint = os.getenv("AZURE_OCR_ENDPOINT")
        key = os.getenv("AZURE_OCR_KEY")
        if not endpoint or not key:
            raise ValueError("AZURE_OCR_ENDPOINT and AZURE_OCR_KEY are required")
        url = endpoint.rstrip("/") + "/vision/v3.2/ocr"
        response = requests.post(
            url,
            headers={"Ocp-Apim-Subscription-Key": key, "Content-Type": "application/octet-stream"},
            data=path.read_bytes(),
            timeout=60,
        )
        response.raise_for_status()
        lines = []
        for region in response.json().get("regions", []):
            for line in region.get("lines", []):
                lines.append(" ".join(word.get("text", "") for word in line.get("words", [])))
        return "\n".join(lines)

    @classmethod
    def _extract_fields(cls, text: str, extraction_type: str) -> Dict[str, Any]:
        gstins = sorted(set(cls.GSTIN_PATTERN.findall(text.upper())))
        pans = sorted(set(cls.PAN_PATTERN.findall(text.upper())))
        tans = sorted(set(cls.TAN_PATTERN.findall(text.upper())))
        invoice_match = cls.INVOICE_PATTERN.search(text)
        amount_match = cls.AMOUNT_PATTERN.search(text)
        tax_values = {name.upper(): float(value.replace(",", "")) for name, value in cls.TAX_PATTERN.findall(text)}
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        extracted = {
            "document_type": extraction_type,
            "vendor": lines[0] if lines else None,
            "customer": None,
            "gstin": gstins[0] if gstins else None,
            "gstins": gstins,
            "pan": pans[0] if pans else None,
            "tan": tans[0] if tans else None,
            "invoice_number": invoice_match.group(1) if invoice_match else None,
            "amount": float(amount_match.group(1).replace(",", "")) if amount_match else None,
            "tax_values": tax_values,
            "bank_statement": cls._extract_bank_statement(lines) if "bank" in extraction_type.lower() else [],
            "raw_text": text[:10000],
        }
        return extracted

    @staticmethod
    def _extract_bank_statement(lines: List[str]) -> List[Dict[str, Any]]:
        transactions = []
        pattern = re.compile(r"(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4}).*?([0-9,]+\.\d{2})")
        for line in lines:
            match = pattern.search(line)
            if match:
                transactions.append({"date": match.group(1), "amount": float(match.group(2).replace(",", "")), "description": line})
        return transactions

    @classmethod
    def extract_document(cls, db: Session, document_id: int, extraction_type: str, provider: str = "local"):
        document = db.query(Document).get(document_id)
        if not document:
            raise ValueError("Document not found")

        text = cls._read_text(document, provider)
        extracted = cls._extract_fields(text, extraction_type)
        extracted["source_file"] = document.file_name
        extracted["provider"] = provider
        confidence = 90 if any([extracted.get("gstin"), extracted.get("pan"), extracted.get("invoice_number"), extracted.get("amount")]) else 40
        log = ExtractionAuditLog(
            document_id=document_id,
            extraction_type=extraction_type,
            provider=provider,
            status="Extracted",
            extracted_data=extracted,
            confidence_score=confidence,
        )
        document.ocr_status = "Extracted"
        document.extracted_data = extracted
        document.document_number = extracted.get("invoice_number") or extracted.get("pan") or extracted.get("tan") or document.document_number
        db.add(log)
        db.commit()
        db.refresh(log)
        return log


class WorkflowService:
    @staticmethod
    def _conditions_match(conditions: Optional[Dict[str, Any]], context: Dict[str, Any]) -> bool:
        if not conditions:
            return True
        for field, expected in conditions.items():
            actual = context.get(field)
            if isinstance(expected, dict):
                if "gte" in expected and not (actual is not None and actual >= expected["gte"]):
                    return False
                if "lte" in expected and not (actual is not None and actual <= expected["lte"]):
                    return False
                if "eq" in expected and actual != expected["eq"]:
                    return False
                if "in" in expected and actual not in expected["in"]:
                    return False
            elif actual != expected:
                return False
        return True

    @classmethod
    def _next_stage(cls, db: Session, workflow_id: int, current_order: int, context: Dict[str, Any]):
        stages = db.query(WorkflowStage).filter(
            WorkflowStage.workflow_id == workflow_id,
            WorkflowStage.stage_order > current_order,
        ).order_by(WorkflowStage.stage_order).all()
        for stage in stages:
            if cls._conditions_match(stage.conditions, context):
                return stage
        return None

    @staticmethod
    def _create_notification(db: Session, instance: WorkflowInstance, stage: Optional[WorkflowStage], message: str):
        if stage and stage.approver_user_id:
            db.add(WorkflowNotification(
                instance_id=instance.id,
                recipient_user_id=stage.approver_user_id,
                channel="InApp",
                subject=f"Workflow action required: {instance.entity_type}",
                message=message,
            ))

    @classmethod
    def create_workflow(cls, db: Session, company_id: int, workflow_name: str, workflow_type: str, stages: List[Dict[str, Any]]):
        if not stages:
            raise ValueError("Workflow must have at least one stage")
        workflow = ApprovalWorkflow(
            company_id=company_id,
            workflow_name=workflow_name,
            workflow_type=workflow_type,
            steps=stages,
            is_active=True,
        )
        db.add(workflow)
        db.flush()
        for index, stage_data in enumerate(stages, start=1):
            db.add(WorkflowStage(
                workflow_id=workflow.id,
                stage_order=index,
                stage_name=stage_data.get("stage_name", f"Stage {index}"),
                approver_role=stage_data.get("approver_role"),
                approver_user_id=stage_data.get("approver_user_id"),
                conditions=stage_data.get("conditions"),
                escalation_rules=stage_data.get("escalation_rules"),
                is_final_stage=index == len(stages),
            ))
        db.commit()
        db.refresh(workflow)
        return workflow

    @classmethod
    def start_instance(cls, db: Session, company_id: int, workflow_id: int, entity_type: str, entity_id: int, initiated_by: Optional[int], context_data: Optional[Dict[str, Any]] = None):
        context = context_data or {}
        first_stage = None
        for stage in db.query(WorkflowStage).filter(WorkflowStage.workflow_id == workflow_id).order_by(WorkflowStage.stage_order).all():
            if cls._conditions_match(stage.conditions, context):
                first_stage = stage
                break
        if not first_stage:
            raise ValueError("Workflow has no matching stages")
        instance = WorkflowInstance(
            company_id=company_id,
            workflow_id=workflow_id,
            entity_type=entity_type,
            entity_id=entity_id,
            current_stage_id=first_stage.id,
            initiated_by=initiated_by,
            context_data=context,
        )
        db.add(instance)
        db.flush()
        db.add(WorkflowHistory(instance_id=instance.id, stage_id=first_stage.id, action="Started", actor_user_id=initiated_by))
        cls._create_notification(db, instance, first_stage, "Approval request has been assigned.")
        db.commit()
        db.refresh(instance)
        return instance

    @classmethod
    def act_on_instance(cls, db: Session, instance_id: int, action: str, actor_user_id: Optional[int] = None, comments: Optional[str] = None):
        instance = db.query(WorkflowInstance).get(instance_id)
        if not instance:
            raise ValueError("Workflow instance not found")

        current_stage = db.query(WorkflowStage).get(instance.current_stage_id) if instance.current_stage_id else None
        context = instance.context_data or {}
        normalized = action.capitalize()
        if normalized == "Approve" and current_stage and not current_stage.is_final_stage:
            next_stage = cls._next_stage(db, instance.workflow_id, current_stage.stage_order, context)
            instance.current_stage_id = next_stage.id if next_stage else None
            if next_stage:
                instance.status = "Pending"
                cls._create_notification(db, instance, next_stage, "Approval request moved to your stage.")
            else:
                instance.status = "Approved"
                instance.completed_at = datetime.utcnow()
        elif normalized == "Approve":
            instance.status = "Approved"
            instance.completed_at = datetime.utcnow()
        elif normalized == "Reject":
            instance.status = "Rejected"
            instance.completed_at = datetime.utcnow()

        db.add(WorkflowHistory(
            instance_id=instance.id,
            stage_id=current_stage.id if current_stage else None,
            action=normalized,
            actor_user_id=actor_user_id,
            comments=comments,
        ))
        db.commit()
        db.refresh(instance)
        return instance

    @classmethod
    def run_escalations(cls, db: Session, company_id: int):
        pending = db.query(WorkflowInstance).filter(
            WorkflowInstance.company_id == company_id,
            WorkflowInstance.status == "Pending",
        ).all()
        escalated = []
        for instance in pending:
            stage = db.query(WorkflowStage).get(instance.current_stage_id) if instance.current_stage_id else None
            rules = stage.escalation_rules if stage else None
            if not rules:
                continue
            after_hours = rules.get("after_hours")
            escalate_to_user_id = rules.get("to_user_id")
            if after_hours and escalate_to_user_id and instance.initiated_at <= datetime.utcnow() - timedelta(hours=after_hours):
                db.add(WorkflowNotification(
                    instance_id=instance.id,
                    recipient_user_id=escalate_to_user_id,
                    channel="InApp",
                    subject="Workflow escalation",
                    message=f"{instance.entity_type} #{instance.entity_id} is pending beyond SLA.",
                ))
                db.add(WorkflowHistory(instance_id=instance.id, stage_id=instance.current_stage_id, action="Escalated", action_data=rules))
                escalated.append(instance.id)
        db.commit()
        return {"escalated_instances": escalated, "count": len(escalated)}


class ComplianceCenterService:
    @classmethod
    def create_task(cls, db: Session, **payload):
        task = ComplianceTask(**payload)
        db.add(task)
        db.commit()
        db.refresh(task)
        return task

    @classmethod
    def create_notice(cls, db: Session, **payload):
        notice = ComplianceNotice(**payload)
        db.add(notice)
        db.commit()
        db.refresh(notice)
        return notice

    @classmethod
    def dashboard(cls, db: Session, company_id: int):
        tasks = db.query(ComplianceTask).filter(ComplianceTask.company_id == company_id).all()
        notices = db.query(ComplianceNotice).filter(ComplianceNotice.company_id == company_id).all()
        pending = [task for task in tasks if task.status == "Pending"]
        overdue = [task for task in pending if task.due_date and task.due_date < date.today()]
        high_risk = [task for task in tasks if task.risk_level in ("High", "Critical")]
        completed = [task for task in tasks if task.status == "Completed"]
        score = round((len(completed) / len(tasks) * 100), 2) if tasks else 100
        return {
            "company_id": company_id,
            "total_tasks": len(tasks),
            "pending_tasks": len(pending),
            "overdue_tasks": len(overdue),
            "high_risk_tasks": len(high_risk),
            "open_notices": len([notice for notice in notices if notice.status == "Open"]),
            "compliance_score": score,
        }

    @classmethod
    def create_roc_filing(cls, db: Session, company_id: int, filing_type: str, financial_year: str, due_date: Optional[date] = None, documents: Optional[List[Dict[str, Any]]] = None):
        filing = RocFiling(
            company_id=company_id,
            filing_type=filing_type,
            financial_year=financial_year,
            due_date=due_date,
            documents=documents or [],
        )
        db.add(filing)
        if due_date:
            db.add(ComplianceTask(
                company_id=company_id,
                task_name=f"{filing_type} filing",
                compliance_type="ROC",
                entity_type="RocFiling",
                financial_year=financial_year,
                due_date=due_date,
                priority="High",
                risk_level="High",
            ))
        db.commit()
        db.refresh(filing)
        return filing

    @classmethod
    def create_director_kyc(cls, db: Session, **payload):
        record = DirectorKYC(**payload)
        db.add(record)
        db.commit()
        db.refresh(record)
        return record


class IntegrationService:
    CONNECTOR_ENDPOINTS = {
        "GST": "GST_PORTAL_ENDPOINT",
        "Income Tax": "INCOME_TAX_PORTAL_ENDPOINT",
        "MCA": "MCA_PORTAL_ENDPOINT",
        "EPFO": "EPFO_PORTAL_ENDPOINT",
        "ESIC": "ESIC_PORTAL_ENDPOINT",
    }

    @classmethod
    def configure(cls, db: Session, company_id: int, integration_type: str, username: Optional[str] = None, api_key: Optional[str] = None):
        integration = GovernmentIntegration(
            company_id=company_id,
            integration_type=integration_type,
            username=username,
            api_key=api_key,
            is_active=True,
            sync_status="Connected",
        )
        db.add(integration)
        db.commit()
        db.refresh(integration)
        return integration

    @classmethod
    def sync(cls, db: Session, integration_id: int, sync_type: str, direction: str = "Pull", payload: Optional[Dict[str, Any]] = None):
        integration = db.query(GovernmentIntegration).get(integration_id)
        if not integration:
            raise ValueError("Integration not found")
        log = IntegrationSyncLog(
            integration_id=integration_id,
            sync_type=sync_type,
            direction=direction,
            status="Running",
            request_payload=payload or {},
        )
        integration.last_sync = datetime.utcnow()
        integration.sync_status = "Running"
        db.add(log)
        db.flush()
        try:
            response_payload = cls._call_connector(integration, sync_type, direction, payload or {})
            log.status = "Success"
            log.response_payload = response_payload
            log.completed_at = datetime.utcnow()
            integration.sync_status = "Connected"
        except Exception as exc:
            log.status = "Error"
            log.error_message = str(exc)
            log.completed_at = datetime.utcnow()
            integration.sync_status = "Error"
        db.commit()
        db.refresh(log)
        return log

    @classmethod
    def _call_connector(cls, integration: GovernmentIntegration, sync_type: str, direction: str, payload: Dict[str, Any]):
        env_name = cls.CONNECTOR_ENDPOINTS.get(integration.integration_type)
        endpoint = os.getenv(env_name or "")
        if not endpoint:
            raise ValueError(f"{integration.integration_type} connector endpoint is not configured")
        headers = {"Content-Type": "application/json"}
        if integration.api_key:
            headers["Authorization"] = f"Bearer {integration.api_key}"
        response = requests.post(
            endpoint.rstrip("/") + f"/sync/{sync_type}",
            json={"direction": direction, "payload": payload, "username": integration.username},
            headers=headers,
            timeout=60,
        )
        response.raise_for_status()
        return response.json()

    @classmethod
    def retry_failed(cls, db: Session, integration_id: int, max_retries: int = 3):
        failed = db.query(IntegrationSyncLog).filter(
            IntegrationSyncLog.integration_id == integration_id,
            IntegrationSyncLog.status == "Error",
            IntegrationSyncLog.retry_count < max_retries,
        ).order_by(IntegrationSyncLog.started_at).all()
        retried = []
        for log in failed:
            log.retry_count += 1
            db.commit()
            retried.append(cls.sync(db, integration_id, log.sync_type, log.direction, log.request_payload or {}).id)
        return {"retried": retried, "count": len(retried)}

    @classmethod
    def record_webhook(cls, db: Session, company_id: int, provider: str, event_type: str, payload: Dict[str, Any]):
        event = WebhookEvent(company_id=company_id, provider=provider, event_type=event_type, payload=payload)
        db.add(event)
        db.commit()
        db.refresh(event)
        return event


class AIInfrastructureService:
    @staticmethod
    def _call_ai(prompt: str, provider: str = "OpenAI", model: Optional[str] = None) -> str:
        if provider.lower() != "openai":
            raise ValueError(f"{provider} provider is not configured")
        if not settings.OPENAI_API_KEY:
            raise ValueError("OPENAI_API_KEY is required for AI provider calls")
        try:
            from openai import OpenAI
        except ImportError as exc:
            raise ValueError("openai package is required for AI provider calls") from exc
        client = OpenAI(api_key=settings.OPENAI_API_KEY)
        response = client.chat.completions.create(
            model=model or settings.OPENAI_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
        )
        return response.choices[0].message.content or ""

    @staticmethod
    def _store_suggestion(db: Session, company_id: int, suggestion_type: str, title: str, description: str, recommendation: str, risk_score: Optional[float] = None):
        suggestion = AISuggestion(
            company_id=company_id,
            suggestion_type=suggestion_type,
            title=title,
            description=description,
            recommendation=recommendation,
            risk_score=risk_score,
        )
        db.add(suggestion)
        db.commit()
        db.refresh(suggestion)
        return suggestion

    @classmethod
    def create_prompt(cls, db: Session, **payload):
        prompt = PromptTemplate(**payload)
        db.add(prompt)
        db.commit()
        db.refresh(prompt)
        return prompt

    @classmethod
    def add_knowledge(cls, db: Session, **payload):
        item = KnowledgeBase(**payload)
        db.add(item)
        db.commit()
        db.refresh(item)
        return item

    @classmethod
    def start_conversation(cls, db: Session, company_id: int, user_id: Optional[int], context_type: str, title: Optional[str] = None):
        conversation = AIConversation(
            company_id=company_id,
            user_id=user_id,
            context_type=context_type,
            conversation_title=title,
            messages=[],
        )
        db.add(conversation)
        db.commit()
        db.refresh(conversation)
        return conversation

    @classmethod
    def add_message(cls, db: Session, conversation_id: int, role: str, content: str):
        conversation = db.query(AIConversation).get(conversation_id)
        if not conversation:
            raise ValueError("Conversation not found")
        messages = conversation.messages or []
        messages.append({"role": role, "content": content, "timestamp": datetime.utcnow().isoformat()})
        conversation.messages = messages
        conversation.updated_at = datetime.utcnow()
        db.commit()
        db.refresh(conversation)
        return conversation

    @classmethod
    def track_usage(cls, db: Session, **payload):
        usage = AIUsageLog(**payload)
        usage.total_tokens = (usage.prompt_tokens or 0) + (usage.completion_tokens or 0)
        db.add(usage)
        db.commit()
        db.refresh(usage)
        return usage

    @classmethod
    def create_suggestion(cls, db: Session, **payload):
        suggestion = AISuggestion(**payload)
        db.add(suggestion)
        db.commit()
        db.refresh(suggestion)
        return suggestion

    @classmethod
    def create_forecast(cls, db: Session, company_id: int, prediction_type: str, period: str, values: Dict[str, Any]):
        if not values:
            values = cls._forecast_from_metrics(db, company_id, prediction_type)
        prediction = AIPrediction(
            company_id=company_id,
            prediction_type=prediction_type,
            period=period,
            predicted_values=values,
            model_version="statistical-v1",
        )
        db.add(prediction)
        db.commit()
        db.refresh(prediction)
        return prediction

    @staticmethod
    def _forecast_from_metrics(db: Session, company_id: int, prediction_type: str):
        metrics = db.query(FinancialMetric).filter(FinancialMetric.company_id == company_id).order_by(FinancialMetric.metric_date.desc()).limit(12).all()
        field_map = {"Revenue": "revenue", "Expense": "expenses", "CashFlow": "cash_flow", "Profit": "profit"}
        field = field_map.get(prediction_type, prediction_type.lower())
        values = [getattr(metric, field, 0) or 0 for metric in reversed(metrics)]
        if not values:
            return {}
        avg_delta = statistics.mean([values[i] - values[i - 1] for i in range(1, len(values))]) if len(values) > 1 else 0
        current = values[-1]
        return {f"month_{i}": round(current + avg_delta * i, 2) for i in range(1, 7)}

    @classmethod
    def analyze_ledgers(cls, db: Session, company_id: int):
        ledgers = db.query(Ledger).filter(Ledger.company_id == company_id).all()
        results = []
        for ledger in ledgers:
            debit = db.query(func.sum(Transaction.debit_amount)).filter(Transaction.ledger_id == ledger.id).scalar() or 0
            credit = db.query(func.sum(Transaction.credit_amount)).filter(Transaction.ledger_id == ledger.id).scalar() or 0
            balance = (ledger.opening_balance or 0) + debit - credit
            results.append({"ledger_id": ledger.id, "name": ledger.name, "debit": debit, "credit": credit, "balance": balance})
        unusual = [item for item in results if abs(item["balance"]) > 1000000]
        return {"ledger_count": len(ledgers), "ledgers": results, "unusual_balances": unusual}

    @classmethod
    def analyze_vouchers(cls, db: Session, company_id: int):
        vouchers = db.query(Voucher).filter(Voucher.company_id == company_id).all()
        unbalanced = [v.id for v in vouchers if round(v.total_debit or 0, 2) != round(v.total_credit or 0, 2)]
        round_amounts = [v.id for v in vouchers if v.total_debit and int(v.total_debit) % 10000 == 0]
        return {"voucher_count": len(vouchers), "unbalanced_vouchers": unbalanced, "round_amount_vouchers": round_amounts}

    @classmethod
    def analyze_gst(cls, db: Session, company_id: int):
        invoices = db.query(Invoice).filter(Invoice.company_id == company_id).all()
        total_tax = sum(invoice.tax_amount or 0 for invoice in invoices)
        invalid_gstin_customers = []
        for invoice in invoices:
            customer = invoice.customer
            if customer and customer.gstin and not re.match(OCRService.GSTIN_PATTERN, customer.gstin):
                invalid_gstin_customers.append({"customer_id": customer.id, "gstin": customer.gstin})
        return {"invoice_count": len(invoices), "total_gst": round(total_tax, 2), "invalid_gstin_customers": invalid_gstin_customers}

    @classmethod
    def analyze_tds(cls, db: Session, company_id: int):
        tds_transactions = db.query(Transaction).join(Voucher).filter(
            Voucher.company_id == company_id,
            Voucher.description.ilike("%TDS%"),
        ).all()
        return {"tds_transaction_count": len(tds_transactions), "tds_amount": round(sum(t.debit_amount or t.credit_amount or 0 for t in tds_transactions), 2)}

    @classmethod
    def tax_advice(cls, db: Session, company_id: int):
        from app.models.models import IncomeTaxReturn
        itr = db.query(IncomeTaxReturn).filter(IncomeTaxReturn.company_id == company_id).order_by(IncomeTaxReturn.created_at.desc()).first()
        if not itr:
            return {"suggestions": ["Create an income tax return draft to calculate specific deduction opportunities."]}
        suggestions = []
        if (itr.section_80_deductions or 0) < 150000:
            suggestions.append("Section 80 deductions appear below the common 80C ceiling; review eligible investments and payments.")
        if (itr.advance_tax or 0) == 0 and (itr.total_tax or 0) > 10000:
            suggestions.append("Advance tax planning is recommended because tax liability exceeds the usual threshold.")
        if (itr.depreciation or 0) == 0 and (itr.business_income or 0) > 0:
            suggestions.append("Review fixed assets for eligible depreciation claims.")
        return {"return_id": itr.id, "taxable_income": itr.taxable_income, "total_tax": itr.total_tax, "suggestions": suggestions}


class CommunicationService:
    @classmethod
    def record_whatsapp_message(cls, db: Session, **payload):
        message = WhatsAppMessage(**payload)
        if message.attachment_url:
            message.processing_status = "AttachmentReceived"
        elif message.message_text and "gst" in message.message_text.lower():
            message.processing_status = "GSTQuery"
            message.processing_log = {"intent": "GSTQuery", "query": message.message_text}
        else:
            message.processing_status = "Received"
        db.add(message)
        db.commit()
        db.refresh(message)
        return message

    @classmethod
    def parse_voice_command(cls, db: Session, company_id: int, transcript: str, user_id: Optional[int] = None):
        lowered = transcript.lower()
        command_type = "General"
        parsed = {"transcript": transcript}
        if "invoice" in lowered:
            command_type = "Invoice"
        elif "purchase" in lowered or "voucher" in lowered:
            command_type = "Voucher"
        elif "gst" in lowered:
            command_type = "GSTQuery"
        amount_match = re.search(r"(?:rs\.?|inr)?\s*([0-9,]+(?:\.\d{1,2})?)", lowered)
        if amount_match:
            parsed["amount"] = float(amount_match.group(1).replace(",", ""))
        ledger_match = re.search(r"(?:ledger|account)\s+([a-z0-9 \-]+)", lowered)
        if ledger_match:
            parsed["ledger_query"] = ledger_match.group(1).strip()
            ledger = db.query(Ledger).filter(
                Ledger.company_id == company_id,
                Ledger.name.ilike(f"%{parsed['ledger_query']}%"),
            ).first()
            if ledger:
                parsed["ledger_id"] = ledger.id
        command = VoiceCommand(
            company_id=company_id,
            user_id=user_id,
            transcript=transcript,
            command_type=command_type,
            parsed_payload={"intent": command_type, **parsed},
            action_status="Parsed",
        )
        db.add(command)
        db.commit()
        db.refresh(command)
        return command


class SecurityService:
    @classmethod
    def log_event(cls, db: Session, **payload):
        event = SecurityLog(**payload)
        db.add(event)
        db.commit()
        db.refresh(event)
        return event

    @classmethod
    def start_session(cls, db: Session, **payload):
        existing = db.query(UserSession).filter(
            UserSession.user_id == payload["user_id"],
            UserSession.is_active == True,
        ).all()
        for session in existing:
            session.is_active = False
            session.ended_at = datetime.utcnow()
        session = UserSession(**payload)
        db.add(session)
        db.commit()
        db.refresh(session)
        return session

    @classmethod
    def register_device(cls, db: Session, **payload):
        device = db.query(DeviceRegistration).filter(
            DeviceRegistration.user_id == payload["user_id"],
            DeviceRegistration.device_id == payload["device_id"],
        ).first()
        if not device:
            device = DeviceRegistration(**payload)
            db.add(device)
        else:
            for key, value in payload.items():
                setattr(device, key, value)
        device.last_seen = datetime.utcnow()
        db.commit()
        db.refresh(device)
        return device


class SaaSAdminService:
    @classmethod
    def create_subscription(cls, db: Session, **payload):
        subscription = TenantSubscription(**payload)
        db.add(subscription)
        db.commit()
        db.refresh(subscription)
        return subscription

    @classmethod
    def record_usage(cls, db: Session, **payload):
        metric = UsageMetric(**payload)
        db.add(metric)
        db.commit()
        db.refresh(metric)
        return metric

    @classmethod
    def revenue_analytics(cls, db: Session):
        subscriptions = db.query(TenantSubscription).filter(TenantSubscription.status == "Active").all()
        plans = {plan.plan_name: plan for plan in db.query(SaaSPlan).all()}
        monthly_revenue = 0
        annual_revenue = 0
        for subscription in subscriptions:
            plan = plans.get(subscription.plan_name)
            if not plan:
                continue
            seats = subscription.seats or 1
            if subscription.billing_cycle == "Annual":
                annual_revenue += (plan.annual_price or 0) * seats
                monthly_revenue += ((plan.annual_price or 0) / 12) * seats
            else:
                monthly_revenue += (plan.monthly_price or 0) * seats
                annual_revenue += (plan.monthly_price or 0) * 12 * seats
        return {
            "active_subscriptions": len(subscriptions),
            "monthly_recurring_revenue": round(monthly_revenue, 2),
            "annual_recurring_revenue": round(annual_revenue, 2),
        }
